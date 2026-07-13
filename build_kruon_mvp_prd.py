from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "kruon_MVP产品需求说明书_2026-07-11.docx"
KEYFRAME = Path("/tmp/kruon_reference_keyframe.png")
SKILL_ROOT = Path("/Users/izzy/.codex/plugins/cache/openai-primary-runtime/documents/26.709.11516/skills/documents")
sys.path.insert(0, str(SKILL_ROOT / "scripts"))
from table_geometry import apply_table_geometry, column_widths_from_weights


NAVY = "14324B"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "2A7F8E"
INK = "172B3A"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_TEAL = "E8F4F4"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "EAF6EE"
PALE_GOLD = "FFF5D6"
PALE_RED = "FDECEC"
WHITE = "FFFFFF"
GRID = "CBD5E1"
RED = "9B1C1C"
GREEN = "216E39"
GOLD = "7A5A00"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 90, "bottom": 90, "start": 120, "end": 120}


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_run_font(run, *, name="Calibri", east_asia="PingFang SC", size=None,
                 bold=None, italic=None, color=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=GRID, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = tc_borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def set_paragraph_border_bottom(paragraph, color=BLUE, size="16", space="5"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_sep, text, fld_end])
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    title.font.size = Pt(27)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = rgb(MUTED)
    subtitle.paragraph_format.space_after = Pt(18)

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        s = styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = rgb(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)


def add_numbering_definition(doc, *, bullet=False):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "PingFang SC")
    r_pr.append(r_fonts)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, numid])


def add_bullets(doc, items, bullet_num_id, *, color=INK, bold_prefix=None):
    for item in items:
        p = doc.add_paragraph()
        apply_num(p, bullet_num_id)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        if bold_prefix and "：" in item:
            prefix, rest = item.split("：", 1)
            r = p.add_run(prefix + "：")
            set_run_font(r, bold=True, color=color)
            r = p.add_run(rest)
            set_run_font(r, color=color)
        else:
            r = p.add_run(item)
            set_run_font(r, color=color)
    return p


def add_numbered(doc, items, num_id, *, compact=False):
    numbering = doc.part.numbering_part.element
    source_num = None
    for el in numbering.findall(qn("w:num")):
        if el.get(qn("w:numId")) == str(num_id):
            source_num = el
            break
    if source_num is None:
        raise ValueError(f"numbering definition not found: {num_id}")
    source_abs = source_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    list_num_id = max(existing_ids, default=0) + 1
    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(list_num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), source_abs)
    new_num.append(abs_id)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    new_num.append(override)
    numbering.append(new_num)
    for item in items:
        p = doc.add_paragraph()
        apply_num(p, list_num_id)
        p.paragraph_format.space_after = Pt(2 if compact else 4)
        p.paragraph_format.line_spacing = 1.10 if compact else 1.25
        r = p.add_run(item)
        set_run_font(r, size=10 if compact else None)


def add_callout(doc, label, text, *, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell, color=fill, size="2")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(label + "  ")
    set_run_font(r, bold=True, color=accent)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    apply_table_geometry(table, [PAGE_WIDTH_DXA], table_width_dxa=PAGE_WIDTH_DXA,
                         indent_dxa=TABLE_INDENT_DXA, cell_margins_dxa=CELL_MARGINS)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, weights, *, header_fill=LIGHT_BLUE,
              font_size=9.2, first_col_bold=False, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, header_fill)
        set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(header)
        set_run_font(r, size=font_size, bold=True, color=NAVY)
    for row in rows:
        new_row = table.add_row()
        set_row_cant_split(new_row)
        cells = new_row.cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.12
            if alignments and i < len(alignments):
                p.alignment = alignments[i]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, bold=(first_col_bold and i == 0), color=INK)
    widths = column_widths_from_weights(weights, PAGE_WIDTH_DXA)
    apply_table_geometry(table, widths, table_width_dxa=PAGE_WIDTH_DXA,
                         indent_dxa=TABLE_INDENT_DXA, cell_margins_dxa=CELL_MARGINS)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    return table


def add_kv_lines(doc, rows):
    for label, value in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(label + "：")
        set_run_font(r, bold=True, color=DARK_BLUE)
        r = p.add_run(value)
        set_run_font(r)


def add_requirement(doc, req_id, title, priority, story, behavior, acceptance,
                    bullet_num_id, *, exception=None):
    p = doc.add_paragraph(style="Heading 3")
    r = p.add_run(f"{req_id}  {title}")
    set_run_font(r, size=12, bold=True, color=DARK_BLUE)
    tag = p.add_run(f"   {priority}")
    set_run_font(tag, size=9, bold=True, color=GREEN if priority == "P0" else GOLD)

    p = doc.add_paragraph()
    r = p.add_run("用户故事：")
    set_run_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run(story)
    set_run_font(r)

    p = doc.add_paragraph()
    r = p.add_run("系统行为")
    set_run_font(r, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(3)
    add_bullets(doc, behavior, bullet_num_id, bold_prefix=True)
    if exception:
        p = doc.add_paragraph()
        r = p.add_run("异常与降级：")
        set_run_font(r, bold=True, color=RED)
        r = p.add_run(exception)
        set_run_font(r)
    p = doc.add_paragraph()
    r = p.add_run("验收标准")
    set_run_font(r, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(3)
    add_bullets(doc, acceptance, bullet_num_id)


def add_source(doc, ref_id, title, url, note):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"[{ref_id}] {title}. ")
    set_run_font(r, size=9.5, bold=True, color=NAVY)
    r = p.add_run(url)
    set_run_font(r, size=9.5, color=BLUE)
    r.font.underline = True
    r = p.add_run(f"；{note}")
    set_run_font(r, size=9.5, color=MUTED)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
configure_styles(doc)
bullet_num_id = add_numbering_definition(doc, bullet=True)
decimal_num_id = add_numbering_definition(doc, bullet=False)

# Running header / footer
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hp.paragraph_format.space_after = Pt(0)
r = hp.add_run("kruon  |  MVP 产品需求说明书")
set_run_font(r, size=8.5, bold=True, color=MUTED)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fp.paragraph_format.space_after = Pt(0)
r = fp.add_run("2026-07-11   ·   ")
set_run_font(r, size=8.5, color=MUTED)
add_field(fp, "PAGE")

# Cover / memo masthead
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(2)
r = p.add_run("PRODUCT REQUIREMENTS DOCUMENT")
set_run_font(r, size=10, bold=True, color=TEAL)

p = doc.add_paragraph(style="Title")
p.add_run("kruon MVP 产品需求说明书")
p = doc.add_paragraph(style="Subtitle")
p.add_run("让已有 AI 工具组成一支可启动、可观察、可干预、可验收的小队")

add_kv_lines(doc, [
    ("版本", "v0.1 - 调研收敛稿"),
    ("状态", "待技术 spike 与用户验证"),
    ("首发范围", "macOS · Codex + Claude Code · 本地优先"),
    ("目标周期", "10 周邀请制 Alpha"),
    ("输入依据", "参考视频、项目路线图、品牌 brief、4 份脑暴材料及外部一手资料"),
])

rule = doc.add_paragraph()
rule.paragraph_format.space_before = Pt(8)
rule.paragraph_format.space_after = Pt(12)
set_paragraph_border_bottom(rule, color=TEAL, size="18")

add_callout(
    doc,
    "核心决策",
    "MVP 的产品价值是本地任务控制闭环，不是完整 3D 模拟。所有关键操作必须在 2D 控制台完成；3D 只投影真实 Run/Event 状态，并且可完全关闭。",
    fill=LIGHT_TEAL,
    accent=TEAL,
)

add_table(doc, ["对象", "MVP 承诺"], [
    ("目标用户", "同时使用多款 AI 编程工具的个人开发者与小型技术团队"),
    ("核心任务", "在一个 Workspace 中派发、监督、审批、取消、收集并验收跨工具任务"),
    ("首发适配器", "Codex 与 Claude Code 深适配；通用 CLI 仅实验性降级"),
    ("北极星", "每周通过 kruon 完成并被用户验收的任务数"),
], [1.55, 4.95], first_col_bold=True, font_size=9.4)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("switch your crew on.")
set_run_font(r, size=11, bold=True, italic=True, color=TEAL)

doc.add_page_break()

# Navigation
doc.add_heading("阅读导航", level=1)
add_table(doc, ["章节", "回答的问题"], [
    ("1. 产品结论", "为什么做、做成什么、怎样算成功"),
    ("2. 调研与甄别", "哪些结论有证据，哪些仍是假设"),
    ("3. 用户与场景", "首发服务谁、解决什么高频工作"),
    ("4-6. 范围与体验", "P0/P1、完整链路、界面与状态世界"),
    ("7. 功能需求", "每个模块的系统行为与验收标准"),
    ("8-11. 技术与安全", "状态、实体、适配器、权限与非功能要求"),
    ("12-15. 验证与交付", "指标、10 周计划、风险、Go/No-Go"),
], [1.6, 4.9], first_col_bold=True)

doc.add_heading("1. 产品结论", level=1)
doc.add_heading("1.1 问题定义", level=2)
p = doc.add_paragraph(
    "重度 AI 用户已经同时使用多个 CLI、桌面 Agent、订阅和 API，但任务、上下文、权限请求、日志与产物分散在不同窗口。"
    "当 Agent 能处理更长、更复杂的工作后，核心瓶颈从“模型会不会做”转向“人怎样同时指挥、监督和验收”。OpenAI 对 Codex App 的官方定位也明确指向多 Agent 并行和长任务监督。[R1]"
)
p = doc.add_paragraph(
    "kruon 不试图替代这些 Agent，也不把所有供应商流量代理到一个新网关。它在用户电脑上建立统一的 Workspace、Task、Run、Approval、Artifact 与 Policy，"
    "让用户在同一套可审计流程里使用已有工具。"
)

doc.add_heading("1.2 一句话定位", level=2)
add_callout(doc, "定位语", "kruon 是面向个人创作者与小型团队的本地 AI 工作指挥空间：连接已有工具，统一任务与权限，把真实工作状态变成一支看得见、管得住、验得过的 AI 小队。", fill=LIGHT_BLUE)

doc.add_heading("1.3 MVP 成功画面", level=2)
add_numbered(doc, [
    "用户在 15 分钟内发现并连接本机 Codex 与 Claude Code。",
    "用户选择一个可信工作目录，设置可写范围、命令、网络和预算策略。",
    "用户创建两个带验收标准的任务，分别交给两个执行器并行运行。",
    "kruon 归一化展示计划、工具调用、文件变更、等待审批、阻塞和完成状态。",
    "高风险动作在执行前触发参数绑定审批；用户可批准、拒绝或缩小范围。",
    "用户可以暂停/取消 Run，并看到明确的终态与残留进程检查。",
    "产物、diff 与测试结果进入交付区，用户人工验收或退回修改。",
    "同一状态流可在 2D 控制台和可关闭的 3D 世界中查看。",
], decimal_num_id)

doc.add_heading("1.4 产品原则", level=2)
add_bullets(doc, [
    "控制优先：先保证启动、状态、权限、取消、产物和验收，再增加表现层。",
    "同源状态：2D 与 3D 只消费同一个 Run/Event 数据源。",
    "人类负责授权与验收：模型输出不能直接等于事实、权限或完成。",
    "本地信任：默认不上传 prompt、文件名、项目名、命令、日志或凭据。",
    "能力诚实：适配器只声明已通过兼容测试的能力；不兼容时降级或阻止。",
    "友好但不幼稚：品牌可以 alive，风险状态必须精确、克制、无歧义。",
], bullet_num_id, bold_prefix=True)

doc.add_page_break()

# Research and evidence
doc.add_heading("2. 调研基础与假设甄别", level=1)
doc.add_heading("2.1 调研输入", level=2)
add_bullets(doc, [
    "实际材料：160.6 秒参考视频、25 页项目分析与路线图、kruon 品牌 brief。",
    "内部假设：Claude、DeepSeek、GPT、Gemini 四份脑暴文档；仅作为假设库。",
    "外部一手资料：Codex、Claude Code、ACP、A2A、MCP、Tauri、OWASP，以及 Paperclip、CC Switch、CLAW3D、openclaw-office 官方页面/仓库。",
    "本机核验：Codex CLI 0.144.0-alpha.4 与 Claude Code 2.1.205 的实际 help/能力表。",
], bullet_num_id, bold_prefix=True)

doc.add_heading("2.2 参考视频可复用的体验语言", level=2)
if KEYFRAME.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    shape = r.add_picture(str(KEYFRAME), width=Inches(6.25))
    shape._inline.docPr.set("descr", "参考视频关键帧：等距虚拟办公室、角色状态、左右侧栏与底部输入区")
    shape._inline.docPr.set("title", "kruon 参考视频关键帧")
    cap = doc.add_paragraph(style="Caption")
    cap.add_run("图 1  参考视频关键帧：等距工作世界 + 左侧角色/任务列表 + 右侧详情 + 底部统一输入。仅作体验证据，不等同于 MVP 交付范围。")

add_table(doc, ["视频证据", "MVP 采用", "甄别结果"], [
    ("角色位置与工位表达状态", "固定空间语义映射", "采纳；必须绑定真实 Run 状态"),
    ("左右侧栏 + 底部输入", "2D 精确控制面", "采纳；减少遮挡，主路径以任务为中心"),
    ("气泡、镜头移动、明暗变化", "轻量反馈", "有限采纳；不得遮挡审批与错误"),
    ("大量角色、会议与复杂场景", "群体模拟", "暂缓；无法证明对任务闭环有必要"),
    ("自由布局/装饰经营", "办公室编辑器", "不进入 MVP；竞品已覆盖且工程量高"),
], [1.9, 1.8, 2.8], header_fill=LIGHT_TEAL, font_size=8.8)

doc.add_heading("2.3 竞品与协议带来的边界", level=2)
add_table(doc, ["对象", "已验证能力", "对 kruon 的约束"], [
    ("Codex App", "多 Agent 并行、项目线程、diff 审阅", "“多 Agent 列表”不是差异化；必须跨厂商"),
    ("Paperclip", "BYO Agent、任务、预算、治理、审计", "控制平面已有强竞品；聚焦本机个人工作"),
    ("CC Switch", "配置、本地路由、failover、用量", "MVP 不重造供应商路由"),
    ("CLAW3D / openclaw-office", "3D/2.5D 办公、状态、会议、任务", "3D 不是壁垒；只保留真实状态投影"),
    ("ACP", "stdio、双向 JSON-RPC、流式、权限请求", "优先协议层，减少脆弱文本解析"),
    ("A2A", "远端 Agent、Task、Artifact、流式与通知", "后续远端接入；不增加 MVP 服务化"),
    ("MCP", "Agent 连接工具/资源与授权", "不是任务编排；禁止 token passthrough"),
], [1.55, 2.2, 2.75], font_size=8.6)

p = doc.add_heading("2.4 脑暴假设判定", level=2)
p.paragraph_format.page_break_before = True
add_table(doc, ["假设", "判定", "MVP 处理"], [
    ("本地优先、BYO 工具", "采纳", "作为首发定位与架构边界"),
    ("统一任务包与人工验收", "采纳", "P0 核心模型"),
    ("按风险而非额度路由", "采纳", "P0 手动选择 + 能力提示"),
    ("共享项目记忆", "需验证", "P1 提案式简报/决策/约束；不做全文同步"),
    ("自动智能路由/fallback", "暂缓", "Alpha 只给建议，不自动换执行器"),
    ("LiteLLM 统一所有订阅", "不采纳", "不接管网页订阅与非官方接口"),
    ("跨设备生产线", "需验证", "保留工作流模板概念；不做同步"),
    ("群聊脑暴", "需验证", "P1：独立产出 -> 质疑 -> 综合"),
    ("广告为主收入", "未证实", "MVP 无广告；商业不得影响路由"),
    ("完整 3D 办公经营", "不采纳", "只做固定微缩场景与状态映射"),
], [2.3, 1.0, 3.2], font_size=8.7, first_col_bold=True)

# Users
doc.add_heading("3. 用户、JTBD 与核心场景", level=1)
doc.add_heading("3.1 首发用户", level=2)
add_callout(doc, "Primary Persona", "AI 重度个人开发者：每天使用至少 2 款 Agent/CLI，在一个或多个本地仓库中完成设计、实现、测试和审查；愿意配置工具，但不愿持续搬运上下文和盯多个终端。", fill=LIGHT_TEAL, accent=TEAL)
add_kv_lines(doc, [
    ("典型环境", "macOS；Git 仓库；Codex + Claude Code；可选 IDE/终端并存"),
    ("核心动机", "减少切窗、等待与重复说明；在高风险动作前保留控制权"),
    ("关键焦虑", "任务卡住无人知、Agent 越权、产物散落、两个 Agent 修改冲突"),
    ("采用阻力", "不信任第三方读取凭据/代码；不想学习复杂编排语言"),
])

doc.add_heading("3.2 暂不作为首发目标", level=2)
add_bullets(doc, [
    "完全不使用 CLI、主要需要通用桌面代办的非技术用户。",
    "需要多租户、SSO、组织预算与合规策略的企业团队。",
    "希望运行 24/7 自治公司的用户；该场景更接近 Paperclip。",
    "把装饰、养成与社交当作主要目的的模拟经营玩家。",
], bullet_num_id)

doc.add_heading("3.3 核心 JTBD", level=2)
add_table(doc, ["触发情境", "我想要", "成功结果"], [
    ("同时有多个开发任务", "把任务交给不同工具并保持边界", "并行不冲突，状态随时可见"),
    ("Agent 长时间运行", "知道正在做什么、是否需要我", "无需轮询终端，重要事件主动提醒"),
    ("Agent 请求高风险操作", "看懂具体影响并决定", "批准绑定参数，可拒绝/缩小范围"),
    ("任务声称完成", "集中审阅产物与测试", "一键验收、退回或打开原生工具"),
    ("换工具继续同一项目", "复用已确认的项目事实", "减少重复交代，来源可追溯"),
], [1.8, 2.2, 2.5], font_size=9)

doc.add_heading("3.4 核心场景优先级", level=2)
add_bullets(doc, [
    "P0 - 双工具并行开发：两个任务、两个隔离 Run、统一事件与产物。",
    "P0 - 高风险审批：文件范围、命令、网络、外部副作用逐项可见。",
    "P0 - 任务验收与退回：diff/文件/测试/完成报告集中归档。",
    "P0 - 副屏监督：世界视图提供低干扰概览，关键事件切换到 2D。",
    "P1 - 项目记忆复用：用户确认后把摘要/决策/约束注入另一执行器。",
    "P1 - 双 Agent 研究评审：独立答案、质疑、综合并保留分歧。",
], bullet_num_id, bold_prefix=True)

# Goals and scope
doc.add_heading("4. MVP 目标、范围与非目标", level=1)
doc.add_heading("4.1 业务与用户目标", level=2)
add_table(doc, ["目标", "成功定义"], [
    ("激活", "新用户 24 小时内连接 2 个工具并完成首个可验收任务"),
    ("控制", "用户能理解 Run 状态、处理审批、暂停/取消并恢复"),
    ("可靠", "终态一致、事件可追溯、失败原因可解释"),
    ("信任", "核心闭环不上传工作内容；危险绕过默认关闭"),
    ("3D 价值", "世界视图提升概览和分享意愿，但关闭后不损失能力"),
], [1.35, 5.15], first_col_bold=True)

doc.add_heading("4.2 P0 / P1 / 非目标", level=2)
add_table(doc, ["层级", "范围"], [
    ("P0 - 必须", "连接与能力探测；Workspace/Policy；任务创建；Codex/Claude 深适配；事件流；审批；暂停/取消；产物与验收；最小 3D；本地日志"),
    ("P1 - Alpha 后半", "提案式项目记忆；双 Agent 研究-质疑-综合；任务复盘快照；通用 CLI 实验适配"),
    ("明确不做", "自动模型路由、额度代理、网页自动化、跨设备/云同步、远端 Agent、插件市场、广告、多人协作、自由建造、角色养成、移动端"),
], [1.45, 5.05], first_col_bold=True, font_size=9.2)

doc.add_heading("4.3 MVP 发布门槛", level=2)
add_bullets(doc, [
    "两款深适配器启动成功率 >=95%，状态终态一致率 >=99%。",
    "正常 Run 的取消请求在 5 秒内进入 cancelling；10 秒内无响应则明确标记 forced-stop required。",
    "所有高风险动作都出现参数绑定审批，拒绝后不得执行同一动作。",
    "用户可在 15 分钟内完成连接、创建 Workspace 和首个任务。",
    "核心路径在关闭 3D 后完全可用；世界视图与 2D 状态不出现分叉。",
    "Alpha 用户首周中位数完成 >=3 个可验收任务。",
], bullet_num_id)

# Experience and IA
doc.add_heading("5. 产品体验与信息架构", level=1)
doc.add_heading("5.1 双层界面", level=2)
add_table(doc, ["层", "职责", "禁止承担"], [
    ("世界视图", "角色/空间/关系/状态概览、轻量反馈、可分享画面", "精确 diff、完整日志、权限参数、配置"),
    ("控制视图", "任务、Run、事件、审批、产物、测试、连接与设置", "用游戏动画代替风险解释"),
], [1.2, 2.65, 2.65], first_col_bold=True)

doc.add_heading("5.2 一级导航", level=2)
add_table(doc, ["入口", "主要对象", "MVP"], [
    ("世界", "Workspace、AgentConnection、Run 状态", "P0"),
    ("任务", "Task、Run、Event、Approval", "P0"),
    ("交付", "Artifact、diff、测试、完成报告", "P0"),
    ("记忆", "项目简报、决策、约束、写入提案", "P1"),
    ("连接", "Codex、Claude Code、能力、版本、健康", "P0"),
    ("设置", "策略、通知、数据、低功耗、诊断", "P0"),
], [1.1, 4.3, 1.1], font_size=9.2, first_col_bold=True,
          alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER])

p = doc.add_heading("5.3 主任务链路", level=2)
add_numbered(doc, [
    "连接：发现工具 -> 版本/认证/能力检查 -> 用户确认连接。",
    "建项目：选择目录 -> 明确信任 -> 配置写入、命令、网络与预算。",
    "建任务：填写目标、上下文、允许范围、验收、测试与回滚。",
    "派发：系统展示执行器能力/风险提示 -> 用户选择 -> 创建 Run。",
    "执行：流式事件归一化 -> 2D/3D 同步 -> 关键事件通知。",
    "审批：展示动作、参数、影响、来源、风险与有效期 -> 批准/拒绝/缩小。",
    "交付：归集文件、diff、报告、测试、链接和执行摘要。",
    "验收：接受并归档；或退回并创建续跑/新 Run。",
], decimal_num_id, compact=True)

# Functional requirements
doc.add_heading("6. 功能需求", level=1)
doc.add_heading("6.1 功能需求明细", level=2)
add_requirement(doc, "FR-01", "工具发现与连接", "P0",
    "作为新用户，我希望 kruon 自动发现可用工具并解释连接能力，使我无需手写复杂配置。",
    [
        "发现：扫描受支持路径并允许手动指定可执行文件。",
        "验证：显示工具名称、版本、认证状态、工作目录能力和结构化输出能力。",
        "授权：连接前展示 kruon 将执行的本地检查；不得读取或导出上游 token。",
        "能力：为每个版本生成 capability manifest，区分 verified / degraded / unsupported。",
    ],
    [
        "全新环境能发现本机 Codex 与 Claude Code，或给出可操作的未发现原因。",
        "认证失效、版本不兼容、命令不可执行均有不同错误码和恢复建议。",
        "用户可断开连接；断开后 kruon 不再启动该工具，但历史记录保留。",
    ], bullet_num_id,
    exception="无法获得结构化事件时，降级为实验性连接并明确列出缺失能力；不得伪装为深适配。")

add_requirement(doc, "FR-02", "Workspace 与策略", "P0",
    "作为开发者，我希望把一个工作目录变成明确的安全边界，并为不同任务复用策略。",
    [
        "信任：首次绑定目录必须由用户确认该目录及其项目配置可信。",
        "范围：配置只读、指定目录写、全 Workspace 写和工作区外写。",
        "命令：支持禁止、白名单、逐次审批和受限执行。",
        "网络：默认关闭或域名白名单；上传与外部写入单独审批。",
        "预算：限制时间、步骤、并发；若上游可提供成本则同时限制金额/token。",
    ],
    [
        "策略在 Run 启动前固化为快照，并可在详情中查看。",
        "策略收紧立即生效；策略放宽只影响后续动作或新 Run。",
        "Workspace 外写入在任何默认模板下都不能静默发生。",
    ], bullet_num_id)

add_requirement(doc, "FR-03", "任务创建与派发", "P0",
    "作为用户，我希望任务在开始前就具有边界和验收条件，避免 Agent 自由发挥后难以收口。",
    [
        "必填：任务目标、Workspace、允许范围和验收标准。",
        "选填：相关文件/简报、测试要求、风险、时限、预算、回滚方式。",
        "建议：根据执行器 capability 和任务风险给出推荐，但由用户最终选择。",
        "并发：MVP 默认最多 2 个 Running Run；超限进入队列。",
    ],
    [
        "缺少验收标准时不能直接执行，只能保存草稿。",
        "选择不具备所需能力的执行器时阻止或明确降级。",
        "同一 Workspace 的潜在文件冲突在启动前提示。",
    ], bullet_num_id,
    exception="MVP 不自动把任务从一个执行器 fallback 到另一个执行器；失败后必须由用户确认重派。")

add_requirement(doc, "FR-04", "Run 生命周期控制", "P0",
    "作为监督者，我希望能看到每次执行的完整生命周期，并在必要时暂停、取消或重试。",
    [
        "启动：使用固定工作目录、策略快照、任务包与适配器版本。",
        "状态：统一映射 queued / planning / running / waiting_approval / blocked / reviewing / completed / failed / cancelling / cancelled / paused。",
        "控制：支持暂停（上游支持时）、取消、强制终止、重试和从会话继续。",
        "幂等：同一控制命令重复提交不会创建额外进程或重复副作用。",
    ],
    [
        "应用重启后能恢复历史 Run，并识别仍存活/已丢失的进程。",
        "取消后无残留子进程；无法确认时标为 uncertain，而不是 completed。",
        "重试创建新 Run ID，并保留与原 Run 的 parent/retry 关系。",
    ], bullet_num_id)

add_requirement(doc, "FR-05", "事件流与可观察性", "P0",
    "作为用户，我希望不读完整终端也能知道 Agent 在做什么，并能按需查看原始证据。",
    [
        "归一化：记录状态、文本摘要、工具调用、文件变更、测试、审批、错误和产物事件。",
        "分层：默认显示可读摘要；原始输出按需展开，并标记来源。",
        "顺序：每个 Run 事件单调编号，乱序事件可检测并重排。",
        "脱敏：凭据和常见 secret 模式在持久化与复制前处理。",
    ],
    [
        "事件详情可追溯到 adapter、时间、原始载荷引用与风险级别。",
        "长输出不会冻结 UI；10000 条事件仍可搜索和筛选。",
        "3D 世界只消费稳定状态/摘要字段，不能解析原始终端文本。",
    ], bullet_num_id)

add_requirement(doc, "FR-06", "权限与审批中心", "P0",
    "作为 Workspace 所有者，我希望在高风险动作前理解真实影响，并作出可审计决定。",
    [
        "分类：文件、命令、网络、外部副作用、记忆写入和预算提升。",
        "展示：动作、精确参数、来源 Run、影响范围、风险、理由和可替代方案。",
        "决定：批准一次、拒绝、缩小参数；MVP 不提供永久批准高风险动作。",
        "完整性：批准绑定 action hash、参数、Run ID 与有效期。",
    ],
    [
        "参数或目标变化后旧批准失效，必须重新请求。",
        "待审批状态通过系统通知和世界视图高优先级提示。",
        "拒绝后 Agent 收到结构化结果，不能通过改写描述复用旧批准。",
    ], bullet_num_id)

add_requirement(doc, "FR-07", "产物归集与人工验收", "P0",
    "作为用户，我希望完成声明对应可检查的产物和测试证据，而不是一段自我评价。",
    [
        "归集：文件、diff、报告、链接、截图、测试结果和完成摘要。",
        "证据：显示产物来源 Run、生成时间、相对 Workspace 路径和变更状态。",
        "验收：接受、退回修改、标记部分接受、打开原生工具。",
        "完成：只有人工接受后 Task 才进入 accepted；Run completed 不等于 Task accepted。",
    ],
    [
        "无产物的任务可提交文字报告，但必须明确任务类型与验收依据。",
        "退回修改生成续跑说明并保留原产物快照。",
        "验收历史不可被后续 Run 静默覆盖。",
    ], bullet_num_id)

add_requirement(doc, "FR-08", "最小 3D 工作世界", "P0",
    "作为副屏监督者，我希望一眼看到谁在工作、等待我、阻塞或完成，同时能快速进入精确详情。",
    [
        "场景：一个固定低多边形微缩办公室，包含工位、白板/规划区、审批点、维修区、评审区、交付区和休息区。",
        "映射：角色、位置、动作和颜色全部由统一 Run 状态驱动。",
        "交互：点击角色打开同一 Run 详情；点击任务/通知聚焦相关角色。",
        "模式：支持关闭、低功耗、减少动画；3D 未加载时自动进入 2D。",
    ],
    [
        "8 类关键状态映射在 2 秒内与 2D 同步。",
        "世界视图失败不影响连接、审批、取消和验收。",
        "不包含自由建造、角色养成、多人联机、GLB 导入和复杂会议模拟。",
    ], bullet_num_id)

add_requirement(doc, "FR-09", "通知、恢复与诊断", "P0",
    "作为用户，我希望只在必须介入时被打断，并能快速定位连接或执行故障。",
    [
        "通知：待审批、阻塞、失败、完成待验收为默认通知；普通日志不弹窗。",
        "恢复：应用崩溃或重启后对存量 Run 执行进程对账。",
        "诊断：提供适配器版本、能力、最近错误、脱敏事件包和兼容测试结果。",
        "隐私：导出诊断包前展示将包含的字段并二次脱敏。",
    ],
    [
        "用户可按 Workspace/事件类型关闭通知，但高风险审批不可静默自动批准。",
        "诊断包默认不包含 prompt、文件内容、API Key 和绝对用户路径。",
        "未知进程状态以 uncertain/needs_attention 表示。",
    ], bullet_num_id)

add_requirement(doc, "FR-10", "项目记忆提案与复用", "P1",
    "作为跨工具用户，我希望把已确认的项目事实复用于下一任务，同时知道内容从哪里来。",
    [
        "类型：项目简报、决策、约束、术语和偏好；不保存原始推理。",
        "写入：Run 只能提交 memory proposal；用户确认后进入规范库。",
        "来源：记录来源 Artifact/Run、创建者、版本、作用域、敏感级别和过期策略。",
        "注入：启动任务前展示将提供给执行器的记忆视图，用户可移除。",
    ],
    [
        "未经确认的提案不能被另一 Agent 当作事实自动注入。",
        "用户可回滚、过期或限制某条记忆的适配器/任务范围。",
        "冲突记忆并存并提示，不由模型静默合并。",
    ], bullet_num_id)

# State model and screens
doc.add_heading("7. 状态模型、界面与交互细则", level=1)
doc.add_heading("7.1 Run 状态与世界映射", level=2)
add_table(doc, ["Run 状态", "世界表达", "用户含义", "允许动作"], [
    ("Queued", "入口/待命区", "等待并发槽位", "取消、调优先级"),
    ("Planning", "白板区", "拆解任务，尚未产生外部副作用", "查看、取消"),
    ("Running", "工位工作", "正在调用工具或修改内容", "查看、暂停/取消"),
    ("Waiting approval", "审批点 + 高亮", "必须由用户决定", "批准、拒绝、缩小"),
    ("Blocked", "维修区", "依赖、认证、权限或测试失败", "诊断、重试、取消"),
    ("Reviewing", "评审区", "检查产物或另一 Run", "查看证据、取消"),
    ("Completed", "交付区", "Run 已结束，Task 等待验收", "验收、退回"),
    ("Paused", "休息区", "执行暂停，资源应释放", "继续、取消"),
    ("Failed/Cancelled", "离线标记", "终止且不可伪装为完成", "诊断、重试"),
], [1.25, 1.55, 2.25, 1.45], font_size=8.0)

doc.add_heading("7.2 关键页面", level=2)
add_table(doc, ["页面", "核心区域", "主操作"], [
    ("首次启动", "价值说明、隐私承诺、工具发现", "连接工具、继续 2D"),
    ("世界", "3D 场景、紧急事件条、任务概览", "聚焦角色、打开详情"),
    ("任务板", "草稿/队列/运行/待验收/完成", "建任务、派发、筛选"),
    ("Run 详情", "状态、时间线、原始事件、策略、控制", "暂停/取消/重试"),
    ("审批中心", "动作参数、影响、来源、风险、有效期", "批准/拒绝/缩小"),
    ("交付中心", "Artifact、diff、测试、摘要", "验收、退回、打开"),
    ("连接管理", "版本、认证、能力、健康、兼容级别", "重测、断开、诊断"),
    ("设置", "策略模板、通知、数据、低功耗", "导出/删除本地数据"),
], [1.35, 3.45, 1.7], font_size=8.8, first_col_bold=True)

doc.add_heading("7.3 待审批交互规范", level=2)
add_bullets(doc, [
    "风险级别不得只用颜色表达；必须有文字标签和影响说明。",
    "默认焦点在“拒绝”或“返回”，不得把批准做成误触主按钮。",
    "外部上传、工作区外写、远端删除和凭据访问必须逐次批准。",
    "审批倒计时只表示有效期，不得制造强迫感；过期后安全失败。",
    "3D 中只显示“需要你决定”，完整参数必须在 2D 审批中心查看。",
], bullet_num_id)

doc.add_heading("7.4 错误与空状态", level=2)
add_bullets(doc, [
    "未连接工具：说明为什么需要连接，并允许只浏览 Demo 事件流。",
    "无任务：提供一个可在示例仓库中运行的安全只读任务模板。",
    "适配器降级：列出缺失能力，例如“无法获取权限事件/取消不可靠”。",
    "3D 不可用：直接落到 2D，不显示技术栈错误给普通用户。",
    "状态不确定：使用 needs_attention，提供进程检查与打开原生工具。",
], bullet_num_id, bold_prefix=True)

# Data and architecture
doc.add_heading("8. 数据模型与事件契约", level=1)
doc.add_heading("8.1 核心实体", level=2)
add_table(doc, ["实体", "职责", "关键关系"], [
    ("Workspace", "项目目录、资料、策略和记忆边界", "1:N Task/Policy/Artifact"),
    ("AgentConnection", "某个具体 CLI 身份与能力", "N:M Workspace/Role"),
    ("Task", "用户意图、范围、验收、预算和时限", "1:N Run"),
    ("Run", "一次可观察执行生命周期", "1:N Event/Artifact/Approval"),
    ("Event", "有序的状态、工具、文件、测试和错误事实", "N:1 Run"),
    ("Approval", "绑定参数与有效期的人类决定", "N:1 Run/Event"),
    ("Artifact", "文件、diff、报告、链接、测试与截图", "N:1 Run/Task"),
    ("Policy", "文件、命令、网络、预算和审批规则", "Snapshot on Run"),
    ("MemoryCapsule", "带来源/作用域/版本的确认事实", "N:M Workspace/AgentView"),
], [1.55, 3.0, 1.95], font_size=8.6)

doc.add_heading("8.2 最小 Event 字段", level=2)
add_table(doc, ["字段", "说明", "要求"], [
    ("event_id / seq", "全局 ID 与 Run 内单调序号", "必填、不可重用"),
    ("run_id / task_id", "执行与任务归属", "必填"),
    ("adapter_id / version", "来源适配器与版本", "必填"),
    ("type / timestamp", "事件类型与发生时间", "必填"),
    ("status_before/after", "状态变更", "状态事件必填"),
    ("summary", "可展示的短摘要", "不得含 secret"),
    ("risk_level", "none/low/medium/high/critical", "动作事件必填"),
    ("payload_ref", "本地原始载荷引用", "按保留策略"),
    ("artifact_refs", "相关产物 ID", "可选"),
], [1.55, 3.15, 1.8], font_size=8.3)

doc.add_heading("8.3 适配器能力分级", level=2)
add_table(doc, ["级别", "接口", "可承诺能力", "MVP 标签"], [
    ("L1", "ACP/官方 SDK/稳定 RPC", "会话、流式、权限、产物、取消", "首选"),
    ("L2", "官方非交互 CLI + JSON/JSONL", "运行、日志、结果、部分控制", "深适配"),
    ("L3", "受控 PTY + 文本解析", "输入输出与粗状态", "实验性"),
    ("L4", "启动外部 App/终端", "导航与上下文", "保底"),
], [0.65, 2.1, 2.7, 1.05], font_size=8.7,
          alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER])

doc.add_page_break()

doc.add_heading("9. 技术架构与首发适配", level=1)
doc.add_heading("9.1 推荐架构", level=2)
add_table(doc, ["模块", "职责", "MVP 技术方向"], [
    ("Desktop Shell", "窗口、托盘、通知、更新、低功耗", "Tauri 2"),
    ("Control Plane", "Task/Run 状态机、队列、取消与恢复", "Rust service + typed commands"),
    ("Adapter Host", "进程启动、协议、版本与能力降级", "独立受限子进程"),
    ("Workspace Service", "目录信任、文件边界、worktree 策略", "本地服务"),
    ("Policy Engine", "文件/命令/网络/预算/审批", "执行前强制校验"),
    ("Observability", "事件、日志、错误、脱敏、诊断", "SQLite + 本地事件总线"),
    ("World Renderer", "事件到 3D 场景投影", "React Three Fiber；无执行权限"),
    ("Memory Hub", "提案、来源、作用域和注入视图", "P1；Markdown 镜像 + SQLite"),
], [1.55, 2.85, 2.1], font_size=8.6)

doc.add_heading("9.2 安全进程边界", level=2)
add_bullets(doc, [
    "3D/展示 WebView 只获得读取稳定状态和发起导航请求的 capability，不得直接启动进程、写文件或读取 Keychain。",
    "所有 Agent 进程由 Adapter Host 启动；环境变量采用 allowlist，不继承无关 secret。",
    "Policy Engine 在模型请求与系统调用之间执行确定性授权，不能依赖模型自报风险。",
    "原始日志与敏感载荷默认本地保存；UI 只通过脱敏摘要访问。",
    "能力配置按窗口拆分，避免 Tauri 多 capability 合并造成意外权限扩大。[R8]",
], bullet_num_id)

doc.add_heading("9.3 Codex 深适配验收 spike", level=2)
add_bullets(doc, [
    "验证 `codex exec --json` 事件覆盖：计划、工具、文件变更、审批、完成与错误。",
    "验证 `--sandbox`、`--cd`、`--add-dir` 与 kruon Policy 的映射。",
    "验证取消信号、子进程清理、应用重启后的 Run 对账。",
    "验证 `resume`、review、output schema 和 diff/测试产物提取。",
    "确认危险绕过审批/沙箱与 hook 信任参数永不由默认 UI 开启。",
], bullet_num_id)

doc.add_heading("9.4 Claude Code 深适配验收 spike", level=2)
add_bullets(doc, [
    "使用 `--output-format stream-json` 与 hook events 建立结构化事件映射。[R2]",
    "验证 allowed/disallowed tools、permission mode、max budget 与 Workspace Policy 的映射。",
    "在 kruon 侧补足非交互模式跳过 Workspace trust 的风险确认。",
    "验证 resume/session id、worktree、后台 Agent、取消与产物识别。",
    "禁止默认使用 bypassPermissions；任何调试入口都必须显著警示并记录。",
], bullet_num_id)

doc.add_heading("10. 安全、隐私与合规需求", level=1)
add_callout(doc, "安全定位", "本地运行不等于天然安全。kruon 同时接触文件、终端、网络、凭据和多个 Agent，安全能力必须出现在主流程，而不是隐藏在设置页。", fill=PALE_RED, accent=RED)

doc.add_heading("10.1 最低控制", level=2)
add_table(doc, ["威胁", "最低控制", "验收证据"], [
    ("间接提示注入", "外部内容不可信标记；动作经策略验证", "恶意文档不能绕过工具策略"),
    ("过度授权", "逐连接 capability + Workspace Policy", "默认最小权限；越界被拒"),
    ("审批操纵", "参数 hash、Run 绑定、过期", "变参后必须重新审批"),
    ("记忆污染", "提案式写入、来源、回滚、过期", "未确认内容不跨 Run"),
    ("凭据泄露", "OS Keychain、环境 allowlist、日志脱敏", "诊断包 secret 扫描为零"),
    ("跨 Agent 放大", "Agent 输出视为不可信输入", "下游不得继承上游权限"),
    ("成本/循环失控", "时间、步骤、并发和预算硬上限", "达到上限安全停止"),
    ("适配器供应链", "版本锁、签名、权限声明、兼容测试", "未知版本自动降级"),
], [1.45, 3.0, 2.05], font_size=8.45)

doc.add_heading("10.2 数据与遥测", level=2)
add_bullets(doc, [
    "默认本地：Task、Run、Event、Artifact 索引、策略与诊断日志。",
    "Keychain：API Key、刷新令牌和 kruon 自有凭据；不得写入 SQLite、记忆或日志。",
    "MVP 无内容遥测：不上传 prompt、模型回复、文件名、绝对路径、项目名、命令、diff 和凭据。",
    "允许的可选匿名事件仅限应用版本、平台、适配器是否连接、功能开关和崩溃类别；默认关闭或首次明确选择。",
    "用户可查看、导出和删除本地数据；删除前说明会保留哪些外部 Agent 会话。",
], bullet_num_id, bold_prefix=True)

doc.add_heading("10.3 安全验收", level=2)
add_bullets(doc, [
    "针对 prompt override、工具滥用、权限提升、记忆污染、数据外泄、递归调用和审批绕过建立回归用例。[R7]",
    "高风险策略、审批逻辑、适配器权限或记忆写入规则变化时，安全测试为发布门。",
    "安全事件记录版本、适配器、策略和结果，但不记录 secret/原始敏感内容。",
    "Alpha 前完成一次威胁建模和人工红队；Critical 问题未关闭不得发布。",
], bullet_num_id)

doc.add_heading("11. 非功能需求", level=1)
add_table(doc, ["类别", "要求", "Alpha 门槛"], [
    ("可靠性", "Run 状态机可恢复；控制命令幂等；错误可解释", "终态一致率 >=99%"),
    ("性能", "10000 条事件可滚动/筛选；3D 不阻塞控制面", "交互 P95 <150ms"),
    ("资源", "2D 常驻低占用；3D 支持低功耗和关闭", "空闲 CPU <2%；3D 关闭显著降耗"),
    ("启动", "冷启动和工具健康检查可分阶段完成", "首屏 <3s；完整检查 <10s"),
    ("兼容", "版本探测 + fixture + 脱敏事件回放", "每款深适配 8 类核心 fixture"),
    ("可访问", "键盘可完成核心路径；状态不只靠颜色", "审批/取消/验收全键盘可达"),
    ("离线", "历史、任务草稿、3D Demo 可离线；执行按上游要求", "无网不丢本地数据"),
    ("更新", "签名更新、可回滚、适配器兼容检查", "失败更新不破坏本地数据库"),
], [1.1, 3.75, 1.65], font_size=8.6, first_col_bold=True)

doc.add_heading("12. 指标、埋点与验证", level=1)
p = doc.add_heading("12.1 北极星与漏斗", level=2)
p.paragraph_format.page_break_before = True
add_callout(doc, "北极星", "每周通过 kruon 完成并被用户人工验收的任务数（Accepted Tasks / WAU）。Run completed 不计入，只有 Task accepted 才计入。", fill=PALE_GREEN, accent=GREEN)
add_table(doc, ["层级", "指标", "用途"], [
    ("激活", "24h 内连接 2 工具并完成首个 accepted Task", "首日价值"),
    ("核心使用", "Accepted Tasks/WAU、跨工具任务占比", "是否成为控制入口"),
    ("可靠性", "启动/取消成功率、终态一致率、事件缺失率", "适配器质量"),
    ("效率", "等待审批时间、人工轮询次数、退回率", "是否减少管理负担"),
    ("留存", "D7/W4、每周活跃 Workspace", "长期价值"),
    ("3D", "主动打开率、停留、点击详情、关闭率", "世界视图是否有用"),
    ("信任", "权限拒绝、遥测关闭、安全投诉/事件", "是否透支信任"),
], [1.15, 3.45, 1.9], font_size=8.7)

doc.add_heading("12.2 Alpha 验证阈值", level=2)
add_table(doc, ["维度", "GO 阈值"], [
    ("问题", "至少 8/12 名受访者每周使用 >=2 个工具，且状态/上下文/权限至少一项为高痛点"),
    ("连接", "70% 目标用户在 15 分钟内连接两个工具"),
    ("任务", "首周中位数 >=3 个 accepted Task"),
    ("可靠", "启动 >=95%；终态一致 >=99%；取消成功 >=95%"),
    ("留存", "邀请制 Alpha W4 >=35%"),
    ("3D", ">=30% 活跃用户每周主动打开，且关闭后无能力损失"),
], [1.3, 5.2], first_col_bold=True, font_size=9.0)

p = doc.add_heading("12.3 调研与测试方法", level=2)
p.paragraph_format.page_break_before = True
add_bullets(doc, [
    "12-20 名目标用户访谈，要求现场展示真实一天的工具切换、等待和上下文搬运。",
    "Codex/Claude 接口 spike：启动、流式、权限、取消、恢复、产物与失败路径。",
    "5-8 名用户完成可用性测试：连接 -> 双任务 -> 审批 -> 验收。",
    "20-50 名邀请 Alpha；先 concierge 支持，不把人工流程伪装成自动化。",
    "事件回放 A/B：相同事件流分别使用 2D-only 与 2D+3D，验证理解速度和主动使用。",
], bullet_num_id)

doc.add_heading("13. 10 周交付计划", level=1)
add_table(doc, ["阶段", "时间", "主要交付", "退出标准"], [
    ("0. 验证", "W1-2", "访谈、工作流录屏、适配器 spike、威胁模型", "确认用户与两款深适配可行"),
    ("1. 控制骨架", "W3-4", "Workspace/Policy、Task/Run、队列、SQLite、2D 任务板", "可启动双工具只读任务"),
    ("2. 闭环", "W5-6", "事件流、审批、取消/恢复、Artifact、验收", "真实任务端到端通过"),
    ("3. 世界视图", "W7-8", "固定 3D 场景、状态映射、低功耗/关闭、通知", "同源状态一致且不阻塞 2D"),
    ("4. Alpha", "W9-10", "诊断、兼容 fixtures、安全回归、打包、onboarding", "20-50 名邀请用户可安装使用"),
], [1.15, 0.85, 2.9, 1.6], font_size=8.35)

doc.add_heading("13.1 团队建议", level=2)
add_bullets(doc, [
    "1 名产品型客户端/全栈：Tauri、React、SQLite、Task/Run 与打包更新。",
    "1 名系统/适配工程师：进程、PTY/协议、权限、取消、跨版本 fixture。",
    "0.5 名 3D/交互：W7 起介入，保持世界渲染与控制平面解耦。",
    "产品/设计由创始人主导，但用户研究、权限 UX 与验收设计不可省略。",
    "Alpha 前安排一次外部安全评审。",
], bullet_num_id, bold_prefix=True)

doc.add_heading("13.2 Definition of Done", level=2)
add_bullets(doc, [
    "需求有可执行验收标准，且自动/手工测试结果已记录。",
    "新增事件、状态、权限或数据字段已更新契约与迁移。",
    "2D 与 3D 使用同一数据；3D 关闭路径完成回归。",
    "日志与诊断包通过 secret/路径脱敏检查。",
    "适配器成功、审批、取消、超时、崩溃、乱码、长输出、版本不兼容 fixture 通过。",
    "用户可理解错误并采取下一步，不以“未知错误”结束。",
], bullet_num_id)

doc.add_heading("14. 风险与应对", level=1)
add_table(doc, ["风险", "概率/影响", "触发信号", "应对"], [
    ("适配器脆弱", "高/高", "版本更新后事件或取消失效", "协议优先、版本锁、fixture、快速降级"),
    ("退化为启动器", "高/高", "用户频繁回原生 App 才能控制", "以权限/取消/产物闭环为发布门"),
    ("3D 成为噱头", "高/中", "只看不点、主动关闭率高", "2D 优先；用主动使用率验证"),
    ("安全事故", "中/极高", "越界写入、泄露、审批绕过", "最小权限、参数审批、红队、fail closed"),
    ("状态不一致", "中/高", "3D/2D/进程不同步", "单一事件源、对账、uncertain 状态"),
    ("范围膨胀", "高/高", "新增工具/平台/市场挤占闭环", "严格 P0/P1，10 周后再决策"),
    ("平台方吸收", "高/高", "原生 App 增加跨 Agent 管理", "跨厂商、本地策略与统一验收"),
    ("记忆污染", "中/高", "错误结论被跨任务复用", "提案式写入、来源、冲突、回滚"),
], [1.25, 1.0, 2.15, 2.1], font_size=8.1)

doc.add_heading("15. Go / Pivot / No-Go 与待决策", level=1)
doc.add_heading("15.1 决策门", level=2)
add_table(doc, ["判断", "标准"], [
    ("GO", "目标用户真实存在高频监督/权限/验收痛点；双工具闭环可稳定实现；第二周仍主动回来使用"),
    ("PIVOT", "用户只需要配置切换与用量 -> 转向 CC Switch 增强；或只需要单 Agent 桌面执行 -> 收窄为单工具监督"),
    ("NO-GO", "核心适配长期只能依赖脆弱文本/界面解析；用户不接受本地权限模型；控制闭环不能显著减少回原生工具"),
], [1.2, 5.3], first_col_bold=True, font_size=9.0)

doc.add_heading("15.2 Alpha 前必须回答", level=2)
add_numbered(doc, [
    "首批用户最常同时使用的两款工具是否确为 Codex 与 Claude Code？",
    "两个适配器能否稳定做到取消、权限参数绑定和产物识别？",
    "并行 Run 是否必须默认使用 worktree，还是允许用户选择隔离策略？",
    "用户最需要的记忆是简报、决策、约束还是历史产物？",
    "3D 在副屏监督中提升了什么：理解速度、介入效率还是分享意愿？",
    "macOS 首发是否足以招募 20-50 名目标用户？",
    "开源边界如何设计，才能让本地控制核心可审计且不妨碍商业持续性？",
], decimal_num_id)

doc.add_heading("15.3 本 PRD 的明确假设", level=2)
add_bullets(doc, [
    "假设 A：用户愿意把工作目录信任交给一个开源或可审计的本地控制层。",
    "假设 B：结构化 CLI 能提供足够稳定的任务生命周期，不必侵入上游 UI。",
    "假设 C：人工验收是高质量多 Agent 工作的必要终点，而不是额外负担。",
    "假设 D：3D 的价值来自真实状态理解和品牌传播，而不是模拟经营深度。",
    "假设 E：跨工具统一任务/权限/产物比自动模型路由更能产生首周复用。",
], bullet_num_id, bold_prefix=True)

doc.add_page_break()

doc.add_heading("附录 A：参考资料与溯源", level=1)
p = doc.add_paragraph(
    "下列外部资料用于验证当前产品与协议能力；功能、版本和市场状态可能变化，进入开发前应再次核验。内部脑暴文档未作为事实来源。"
)
add_source(doc, "R1", "OpenAI - Introducing the Codex app", "https://openai.com/index/introducing-the-codex-app/", "多 Agent 并行、项目线程与长任务监督")
add_source(doc, "R2", "Claude Code - CLI reference", "https://code.claude.com/docs/en/cli-usage", "stream-json、会话恢复、权限与 hook 事件")
add_source(doc, "R3", "Agent Client Protocol - Architecture", "https://agentclientprotocol.com/get-started/architecture", "stdio、双向 JSON-RPC、流式与权限请求")
add_source(doc, "R4", "A2A Protocol Specification", "https://github.com/a2aproject/A2A/blob/main/docs/specification.md", "Task、Artifact、Agent Card 与生命周期")
add_source(doc, "R5", "Model Context Protocol - Authorization", "https://modelcontextprotocol.io/specification/2025-06-18/basic/authorization", "token audience 绑定与禁止透传")
add_source(doc, "R6", "OWASP - MCP Security Cheat Sheet", "https://cheatsheetseries.owasp.org/cheatsheets/MCP_Security_Cheat_Sheet.html", "MCP 工具与供应链风险")
add_source(doc, "R7", "OWASP - AI Agent Security Cheat Sheet", "https://cheatsheetseries.owasp.org/cheatsheets/AI_Agent_Security_Cheat_Sheet.html", "最小权限、HITL、记忆隔离与安全测试")
add_source(doc, "R8", "Tauri 2 - Capabilities", "https://v2.tauri.app/security/capabilities/", "按窗口/WebView 限制系统能力")
add_source(doc, "R9", "Paperclip", "https://github.com/paperclipai/paperclip", "BYO Agent、任务、预算、治理与审计控制平面")
add_source(doc, "R10", "CC Switch", "https://ccswitch.co/", "多工具配置、本地路由、failover、用量与本地密钥")
add_source(doc, "R11", "OpenClaw Office", "https://github.com/WW-AI-Lab/openclaw-office/blob/main/README.en.md", "2D 等距数字孪生办公室与状态/协作可视化")
add_source(doc, "R12", "CLAW3D", "https://claw3d.site/", "3D 办公、任务、PR、会议、跟随与布局编辑")

doc.add_heading("附录 B：内部材料", level=1)
add_bullets(doc, [
    "参考.mp4 - 160.6 秒 3D/2.5D 虚拟办公室产品演示。",
    "kruon_项目分析与产品路线图_2026-07-11.docx - 产品战略、技术、安全、商业与路线图。",
    "brand/kruon-brand-brief-v0.1.md - 品牌名称、定位、口号与关键词。",
    "brainstorm/Claude-AI控制中心项目分析.md - 竞品、定位、适配器、记忆、技术与商业假设。",
    "brainstorm/DeepSeek-LiteLLM思路.md - LiteLLM 网关、自动路由与创意流水线假设。",
    "brainstorm/GPT-已有项目如何管理ai工具.md - 仓库真相源、任务包、风险路由、验收与工作室流程。",
    "brainstorm/Gemini-虚拟办公室形态Multi-Agent模拟器.md - 3D 工作室、角色映射与商业化脑暴。",
], bullet_num_id)

doc.add_heading("附录 C：术语", level=1)
add_table(doc, ["术语", "定义"], [
    ("AgentConnection", "一个可调用的具体工具身份，而不是抽象模型名称。"),
    ("Task", "包含目标、边界与验收标准的用户工作单元。"),
    ("Run", "Task 在某个 AgentConnection 上的一次执行生命周期。"),
    ("Accepted", "用户已检查产物并明确接受；区别于 Run Completed。"),
    ("Capability Manifest", "适配器在特定版本上已验证可承诺的能力清单。"),
    ("Memory Proposal", "Agent 提交、但尚未被用户确认的可复用事实/决策。"),
    ("World View", "由统一事件流驱动的 3D 状态投影，不是执行或权限边界。"),
], [1.85, 4.65], first_col_bold=True, font_size=9.0)

# Core properties
props = doc.core_properties
props.title = "kruon MVP 产品需求说明书"
props.subject = "本地 AI 工作指挥空间 MVP PRD"
props.author = "kruon 项目组"
props.keywords = "kruon, MVP, PRD, AI Agent, local control plane, 3D workspace"
props.comments = "基于项目路线图、参考视频、品牌资料、脑暴甄别与外部调研形成。"

doc.save(OUTPUT)
print(OUTPUT)
